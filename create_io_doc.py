#!/usr/bin/env python
"""Generate Word document: Closed-Loop Evaluation Inputs & Outputs."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ---- Page setup ----
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# ---- Styles ----
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elm)

def add_styled_table(doc, headers, rows, col_widths=None, header_color="1565C0"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if r_idx % 2 == 0:
                set_cell_shading(cell, "F5F5F5")

    # Set column widths
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    return table

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph("")
doc.add_paragraph("")
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("End-to-End Autonomous Driving\nClosed-Loop Evaluation")
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(21, 101, 192)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Input / Output Specification Document")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph("")
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("Digital Twin Simulation | CARLA Simulator\nMLOps Pipeline | CI/CD Integration")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(120, 120, 120)

doc.add_paragraph("")
doc.add_paragraph("")
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run("Version 1.0 | April 2026")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(150, 150, 150)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Overview",
    "2. System Architecture Summary",
    "3. Sensor Inputs (from CARLA / Vehicle)",
    "4. Derived / Computed Inputs (Pre-processed)",
    "5. Training-Only Inputs (Ground Truth)",
    "6. Model Outputs (E2E Network)",
    "7. Control Outputs (PID Controller)",
    "8. Evaluation Metric Outputs",
    "9. Logged Outputs (for Retraining)",
    "10. Data Flow Summary",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ============================================================
# 1. OVERVIEW
# ============================================================
doc.add_heading("1. Overview", level=1)
doc.add_paragraph(
    "This document specifies all inputs consumed and outputs produced during the closed-loop "
    "evaluation of an End-to-End (E2E) autonomous driving model within the CARLA digital twin "
    "simulator. The evaluation runs in a synchronous simulation loop at 20 Hz, where at each "
    "tick the model receives sensor data, performs inference, and produces vehicle control "
    "commands that are applied to the ego vehicle in the simulated environment."
)
doc.add_paragraph(
    "The evaluation covers 220 predefined driving scenarios (Bench2Drive benchmark) across "
    "multiple towns, weather conditions, and traffic situations. Results are aggregated into "
    "standardized metrics (driving score, collision rate, route completion) that determine "
    "whether a model candidate is safe for deployment."
)

# ============================================================
# 2. SYSTEM ARCHITECTURE SUMMARY
# ============================================================
doc.add_heading("2. System Architecture Summary", level=1)
doc.add_paragraph(
    "The closed-loop evaluation pipeline consists of five stages that execute in a continuous loop:"
)
stages = [
    ("Stage 1: Environment Setup", "Load route definition, initialize CARLA world (town, weather, traffic), spawn ego vehicle, attach sensor array, activate scenario triggers."),
    ("Stage 2: Sensor Data Capture", "At each simulation tick (20 Hz), capture RGB images from 6 surround cameras, LiDAR point cloud, IMU readings (acceleration, angular velocity, compass), GPS position, and vehicle speed."),
    ("Stage 3: E2E Model Inference", "Process sensor inputs through the neural network: image backbone extracts visual features, BEV encoder creates a bird's-eye-view representation, detection head identifies surrounding objects, map head predicts lane structure, trajectory VAE generates candidate paths, and the planning head selects the optimal collision-free trajectory."),
    ("Stage 4: Control & Actuation", "Convert the selected trajectory into low-level vehicle control commands (throttle, steering, brake) using a PID controller. Apply commands to the ego vehicle via the CARLA API and advance the simulation by one tick."),
    ("Stage 5: Metrics & Logging", "Record evaluation metrics (driving score, collisions, lateral error, route completion, speed profile). Log all sensor data, predictions, and control commands for later analysis and model retraining."),
]
for title_text, desc in stages:
    p = doc.add_paragraph()
    run = p.add_run(title_text + ": ")
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(desc)
    run2.font.size = Pt(10)

doc.add_page_break()

# ============================================================
# 3. SENSOR INPUTS
# ============================================================
doc.add_heading("3. Sensor Inputs (from CARLA / Vehicle)", level=1)
doc.add_paragraph(
    "These are the raw sensor signals captured from the simulated vehicle at each simulation tick. "
    "They mirror the physical sensors that would be mounted on a real autonomous vehicle."
)

sensor_headers = ["#", "Input Signal", "Data Format", "Specifications", "Description"]
sensor_rows = [
    ["1", "CAM_FRONT", "RGB Image\n1600 x 900 px", "FOV: 70 deg\nPos: x=0.80, z=1.60\nYaw: 0 deg", "Forward-facing camera capturing the road ahead, traffic lights, lead vehicles, and road signs."],
    ["2", "CAM_FRONT_LEFT", "RGB Image\n1600 x 900 px", "FOV: 70 deg\nPos: x=0.27, y=-0.55\nYaw: -55 deg", "Left-angled camera covering the front-left blind spot, adjacent lane traffic at intersections."],
    ["3", "CAM_FRONT_RIGHT", "RGB Image\n1600 x 900 px", "FOV: 70 deg\nPos: x=0.27, y=0.55\nYaw: +55 deg", "Right-angled camera covering the front-right blind spot and adjacent lane traffic."],
    ["4", "CAM_BACK", "RGB Image\n1600 x 900 px", "FOV: 110 deg\nPos: x=-2.0, z=1.60\nYaw: 180 deg", "Rear-facing wide-angle camera for reversing, rear traffic monitoring, and lane change safety."],
    ["5", "CAM_BACK_LEFT", "RGB Image\n1600 x 900 px", "FOV: 70 deg\nPos: x=-0.32, y=-0.55\nYaw: -110 deg", "Rear-left camera covering the left-side blind spot during lane changes and turns."],
    ["6", "CAM_BACK_RIGHT", "RGB Image\n1600 x 900 px", "FOV: 70 deg\nPos: x=-0.32, y=0.55\nYaw: +110 deg", "Rear-right camera covering the right-side blind spot during lane changes and merges."],
    ["7", "GPS (GNSS)", "float[2]\n[latitude, longitude]", "Rate: 100 Hz\nPos: x=-1.4, z=0.0", "Global positioning from satellite navigation. Provides absolute world position of the ego vehicle."],
    ["8", "IMU", "float[7]\n[ax,ay,az, gx,gy,gz,\ncompass]", "Rate: 20 Hz\nPos: x=-1.4, z=0.0", "Inertial measurement unit providing 3-axis acceleration (m/s2), 3-axis angular velocity (rad/s), and magnetic compass heading (rad)."],
    ["9", "Speedometer", "float (m/s)", "Rate: 20 Hz", "Current longitudinal speed of the ego vehicle from wheel odometry."],
    ["10", "BEV Camera", "RGB Image\n512 x 512 px", "FOV: 50 deg\nPos: z=50.0 (overhead)\nPitch: -90 deg", "Top-down bird's eye view camera mounted 50m above the vehicle. Used for visualization and debugging only."],
]
add_styled_table(doc, sensor_headers, sensor_rows, col_widths=[1, 3, 3, 3.5, 6.5])

doc.add_page_break()

# ============================================================
# 4. DERIVED INPUTS
# ============================================================
doc.add_heading("4. Derived / Computed Inputs (Pre-processed)", level=1)
doc.add_paragraph(
    "These inputs are computed from the raw sensor data before being fed to the E2E model. "
    "They include coordinate transforms, fused ego state vectors, and navigation commands."
)

derived_headers = ["#", "Input Signal", "Data Format", "Description"]
derived_rows = [
    ["11", "can_bus", "float[18]", "Fused ego state vector containing:\n- [0:2] Position (x, y) in world frame\n- [3:7] Orientation as quaternion\n- [7] Speed (m/s)\n- [10:13] Linear acceleration (x, y, z)\n- [13:16] Angular velocity (x, y, z)\n- [16] Yaw angle (radians)\n- [17] Yaw angle (degrees)"],
    ["12", "ego_lcf_feat", "float[9]", "Ego lateral control features:\n- [0:2] Position (x, y)\n- [2:4] Acceleration (x, y)\n- [4] Rotation (quaternion w)\n- [5] Vehicle length (4.89 m)\n- [6] Vehicle width (1.84 m)\n- [7] Distance from origin\n- [8] Previous steering value"],
    ["13", "ego_fut_cmd", "one-hot[6]", "High-level navigation command encoded as one-hot vector:\n0=Turn Left, 1=Turn Right, 2=Go Straight,\n3=Follow Lane, 4=Change Left, 5=Change Right"],
    ["14", "command_near_xy", "float[2]", "Next target waypoint transformed into the ego vehicle's local coordinate frame (relative x, y position)."],
    ["15", "lidar2img", "float[6 x 4 x 4]", "Projection matrices for each of the 6 cameras. Transforms 3D LiDAR points to 2D image pixel coordinates. Used by BEV encoder for view transformation."],
    ["16", "lidar2cam", "float[6 x 4 x 4]", "Extrinsic transformation matrices from LiDAR frame to each camera frame. Encodes the physical mounting positions and orientations of all cameras."],
    ["17", "l2g_r_mat", "float[3 x 3]", "Rotation matrix from LiDAR coordinate frame to global (world) coordinate frame. Updated each tick based on current ego orientation."],
    ["18", "l2g_t", "float[3]", "Translation vector from LiDAR frame to global frame. Updated each tick based on current ego position."],
]
add_styled_table(doc, derived_headers, derived_rows, col_widths=[1, 3, 3, 10])

doc.add_page_break()

# ============================================================
# 5. TRAINING-ONLY INPUTS
# ============================================================
doc.add_heading("5. Training-Only Inputs (Ground Truth)", level=1)
doc.add_paragraph(
    "These inputs are available only during training (from the dataset annotations). "
    "They are NOT used during closed-loop inference/evaluation. Listed here for completeness."
)

gt_headers = ["#", "Input Signal", "Data Format", "Description"]
gt_rows = [
    ["19", "gt_bboxes_3d", "list of 3D boxes\n[x,y,z,w,l,h,yaw]", "Ground truth 3D bounding boxes for all surrounding agents (vehicles, pedestrians, cyclists). Used as supervision for the detection head during training."],
    ["20", "gt_labels_3d", "int[]", "Class labels for each ground truth bounding box (car, truck, pedestrian, bicycle, etc)."],
    ["21", "gt_attr_labels", "int[]", "Attribute labels for each object: moving, stopped, parked. Provides behavioral context for motion prediction training."],
    ["22", "ego_his_trajs", "float[4 x 2]", "Past 4 ego vehicle positions (x, y) as a history trajectory. Used to condition trajectory prediction on recent motion."],
    ["23", "ego_fut_trajs", "float[6 x 2]", "Ground truth future 6 ego positions (x, y). Used as supervision signal for the trajectory prediction and planning heads."],
    ["24", "ego_fut_masks", "bool[6]", "Validity mask for each future trajectory point. Handles cases where the route ends before 6 future steps."],
]
add_styled_table(doc, gt_headers, gt_rows, col_widths=[1, 3, 3, 10])

doc.add_page_break()

# ============================================================
# 6. MODEL OUTPUTS
# ============================================================
doc.add_heading("6. Model Outputs (E2E Network)", level=1)
doc.add_paragraph(
    "These are the direct outputs of the E2E neural network's forward pass. "
    "The model processes all sensor inputs in a single inference step and produces "
    "perception, prediction, and planning outputs simultaneously."
)

model_headers = ["#", "Output Signal", "Data Format", "Description"]
model_rows = [
    ["1", "ego_fut_preds\n(Trajectory Candidates)", "float[6 x 6 x 2]", "6 candidate future trajectories, one per navigation command mode. Each trajectory has 6 waypoints, each waypoint is (delta_x, delta_y) relative to the current position. The trajectory matching the current navigation command is selected for execution."],
    ["2", "3D Object Detections", "list of\n[x,y,z,w,l,h,yaw,\nvx,vy,class,score]", "Detected surrounding agents with 3D bounding boxes: center position (x,y,z), dimensions (width, length, height), orientation (yaw), velocity (vx, vy), object class, and confidence score."],
    ["3", "Map Predictions", "list of polylines\n+ class labels", "Predicted local road structure: lane center lines, lane boundaries, road edges, crosswalks, and drivable area boundaries as vectorized polylines with semantic labels."],
    ["4", "Trajectory Scores", "float[6]", "Confidence/safety score for each of the 6 candidate trajectories. Incorporates collision risk, lane boundary violation probability, and direction compliance."],
]
add_styled_table(doc, model_headers, model_rows, col_widths=[1, 3.5, 3, 9.5], header_color="B71C1C")

doc.add_page_break()

# ============================================================
# 7. CONTROL OUTPUTS
# ============================================================
doc.add_heading("7. Control Outputs (PID Controller)", level=1)
doc.add_paragraph(
    "The selected trajectory from the model is converted into low-level vehicle control "
    "commands by a PID controller. These commands are applied to the ego vehicle at each tick."
)

ctrl_headers = ["#", "Output Signal", "Data Format", "Range", "Description"]
ctrl_rows = [
    ["5", "Throttle", "float", "0.0 - 0.75", "Acceleration command. Controls how much the vehicle accelerates. Capped at 0.75 to prevent excessive speed. Determined by longitudinal PID controller based on speed error relative to target speed."],
    ["6", "Steering", "float", "-1.0 to +1.0", "Steering angle command. -1.0 = full left turn, 0.0 = straight, +1.0 = full right turn. Determined by lateral PID controller tracking the trajectory waypoints."],
    ["7", "Brake", "float", "0.0 - 1.0", "Braking force command. Applied when the vehicle exceeds target speed or needs to stop. Mutually exclusive with throttle (if throttle > brake, brake is set to 0)."],
]
add_styled_table(doc, ctrl_headers, ctrl_rows, col_widths=[1, 2.5, 2, 2.5, 9], header_color="1B5E20")

doc.add_paragraph("")
doc.add_heading("Control Logic", level=2)
doc.add_paragraph(
    "The PID controller converts the planned trajectory into control commands using two loops:"
)
items = [
    "Lateral Control: Computes cross-track error between the vehicle heading and the direction to the next trajectory waypoint. PID gains: Kp=1.2, Ki=0.0, Kd=0.3.",
    "Longitudinal Control: Computes speed error between current speed and target speed (30 km/h). PID gains: Kp=0.5, Ki=0.05, Kd=0.1.",
    "Safety Rule: If brake < 0.05, brake is set to 0. If throttle > brake, brake is set to 0 (no simultaneous throttle + brake).",
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============================================================
# 8. EVALUATION METRICS
# ============================================================
doc.add_heading("8. Evaluation Metric Outputs", level=1)
doc.add_paragraph(
    "These metrics are computed across the entire route evaluation and aggregated across "
    "all 220 scenarios. They determine whether a model candidate passes the validation gate."
)

metric_headers = ["#", "Metric", "Unit", "Description", "Pass Criteria"]
metric_rows = [
    ["8", "Driving Score", "0 - 100", "Primary composite metric: Route Completion % multiplied by an infraction penalty multiplier. Penalizes collisions, red light violations, and stop sign violations.", "> baseline model"],
    ["9", "Route Completion", "0 - 100 %", "Percentage of route waypoints successfully reached by the ego vehicle before timeout or terminal failure.", "> 80%"],
    ["10", "Collision Count", "integer", "Total number of collisions with other vehicles, pedestrians, cyclists, or static objects during the route.", "= 0 (ideal)"],
    ["11", "Lane Invasions", "integer", "Number of times the vehicle crossed a lane boundary. Includes both opposing lane and sidewalk invasions.", "< 5 per route"],
    ["12", "Lateral Error (avg)", "meters", "Average perpendicular distance between the actual vehicle path and the planned route centerline.", "< 1.0 m"],
    ["13", "Lateral Error (max)", "meters", "Maximum perpendicular distance between the actual path and the planned route. Indicates worst-case deviation.", "< 3.0 m"],
    ["14", "Average Speed", "km/h", "Mean speed across all simulation ticks. Should be close to the speed limit without excessive stops.", "20-40 km/h"],
    ["15", "Success Rate", "0 or 1", "Binary: 1 if the vehicle completed the entire route without any terminal failure (e.g., collision, timeout, stuck).", "> 15% across all routes"],
]
add_styled_table(doc, metric_headers, metric_rows, col_widths=[1, 3, 2, 7, 4], header_color="E65100")

doc.add_page_break()

# ============================================================
# 9. LOGGED OUTPUTS
# ============================================================
doc.add_heading("9. Logged Outputs (for Retraining)", level=1)
doc.add_paragraph(
    "All evaluation data is logged and uploaded to the data lake. This data is used for "
    "model improvement, failure analysis, and continuous retraining via the MLOps pipeline."
)

log_headers = ["#", "Output", "Format", "Frequency", "Description"]
log_rows = [
    ["16", "RGB Camera Frames", "PNG images\n(6 cameras)", "Every 10 ticks\n(2 Hz)", "All 6 surround camera images saved as compressed PNG files. Used for failure replay, annotation, and retraining dataset creation."],
    ["17", "BEV Frame", "PNG image\n512 x 512", "Every 10 ticks\n(2 Hz)", "Top-down bird's eye view image for visualization and debugging of the planned trajectory."],
    ["18", "Metadata JSON", "JSON file\nper frame", "Every 10 ticks\n(2 Hz)", "Per-frame metadata: planned trajectory (6 waypoints), all 6 candidate trajectories, control commands (throttle/steer/brake), navigation command index, PID controller state."],
    ["19", "Step Data JSON", "JSON array\n(full time series)", "Once per route", "Complete time series at 20 Hz: actual position, expected position, lateral error, speed, heading, velocity, acceleration, control commands, collision count, lane invasion count."],
    ["20", "Expected Route JSON", "JSON array\nof waypoints", "Once per route", "Full planned route: all waypoints with (x, y, z), road ID, lane ID. Serves as ground truth reference for trajectory comparison."],
    ["21", "Actual Trajectory JSON", "JSON array\nof positions", "Once per route", "Actual vehicle positions at every tick: (x, y, z) in world coordinates. Compared against expected route for deviation analysis."],
    ["22", "Results Summary JSON", "JSON object", "Once per route", "Aggregated evaluation results: driving score, total distance, average/max speed, average/max lateral error, collision count, lane invasions, route completion percentage."],
]
add_styled_table(doc, log_headers, log_rows, col_widths=[1, 3, 2.5, 2.5, 8], header_color="6A1B9A")

doc.add_page_break()

# ============================================================
# 10. DATA FLOW SUMMARY
# ============================================================
doc.add_heading("10. Data Flow Summary", level=1)
doc.add_paragraph(
    "The following diagram shows how data flows through the closed-loop evaluation system:"
)

flow_items = [
    ("CARLA Simulator", "Produces sensor data (6 cameras, GPS, IMU, speed) at 20 Hz per tick"),
    ("Pre-processing", "Fuses raw sensors into can_bus[18], ego_lcf_feat[9], ego_fut_cmd[6], camera projection matrices"),
    ("E2E Model Inference", "Consumes 6 RGB images + ego state + route command. Produces 6 candidate trajectories + 3D detections + map predictions"),
    ("Trajectory Selection", "Picks the trajectory matching the current navigation command (turn left / go straight / etc)"),
    ("PID Controller", "Converts trajectory waypoints into throttle [0-0.75], steering [-1,+1], brake [0-1]"),
    ("CARLA Vehicle API", "Applies control commands. Vehicle moves. World advances one tick."),
    ("Metrics Logger", "Records position, speed, lateral error, collisions, lane invasions at every tick"),
    ("Data Lake Upload", "Saves all frames, predictions, and metrics for retraining pipeline"),
]

table = doc.add_table(rows=len(flow_items), cols=3)
table.style = 'Table Grid'
for i, (component, description) in enumerate(flow_items):
    table.rows[i].cells[0].text = str(i + 1)
    p = table.rows[i].cells[1].paragraphs[0]
    run = p.add_run(component)
    run.bold = True
    run.font.size = Pt(10)
    table.rows[i].cells[2].text = description
    if i % 2 == 0:
        set_cell_shading(table.rows[i].cells[0], "E3F2FD")
        set_cell_shading(table.rows[i].cells[1], "E3F2FD")
        set_cell_shading(table.rows[i].cells[2], "E3F2FD")

# Summary counts
doc.add_paragraph("")
doc.add_heading("Summary Counts", level=2)
summary_table = doc.add_table(rows=5, cols=2)
summary_table.style = 'Table Grid'
summary_data = [
    ("Total Sensor Inputs", "10 signals (6 cameras + GPS + IMU + Speed + BEV)"),
    ("Total Derived Inputs", "8 signals (can_bus, ego features, commands, transforms)"),
    ("Total Model Outputs", "4 signals (trajectories, detections, map, scores)"),
    ("Total Control Outputs", "3 signals (throttle, steering, brake)"),
    ("Total Evaluation Metrics", "8 metrics (driving score, collisions, lateral error, etc)"),
]
for i, (label, value) in enumerate(summary_data):
    p = summary_table.rows[i].cells[0].paragraphs[0]
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(10)
    summary_table.rows[i].cells[1].text = value
    set_cell_shading(summary_table.rows[i].cells[0], "1565C0")
    p.runs[0].font.color.rgb = RGBColor(255, 255, 255)

# ---- Save ----
output_path = r"C:\GenAD\architecture\Closed_Loop_Evaluation_IO_Specification.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
